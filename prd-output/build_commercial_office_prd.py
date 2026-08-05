from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "prd-output"
OUTPUT = OUT_DIR / "透明房产小程序V1.1.0_商业办公功能需求文档.docx"

FONT = "Microsoft YaHei"
INK = "222222"
MUTED = "697386"
BLUE = "0B5CAD"
LIGHT_BLUE = "EAF3FF"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9DEE7"
WHITE = "FFFFFF"
RED = "C62828"
GREEN = "1D7A46"

PAGE_WIDTH_DXA = 11906  # A4
PAGE_HEIGHT_DXA = 16838
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=10.5, bold=None, color=INK, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_WIDTH_DXA, (widths, sum(widths))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=MID_GRAY, size=4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(str(value))
        set_run_font(run, size=font_size, bold=True)

    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    return table


def add_para(doc, text="", size=10.5, bold=False, color=INK, before=0, after=5, line=1.15, align=None, italic=False, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if keep:
        p.paragraph_format.keep_with_next = True
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(text)
    set_run_font(run, size=10)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=accent, size=6)
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=accent)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(body)
    set_run_font(r, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size={1: 15, 2: 12.5, 3: 11}[level], bold=True, color=INK)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    specs = {
        "Heading 1": (15, 14, 7),
        "Heading 2": (12.5, 10, 5),
        "Heading 3": (11, 7, 3),
    }
    for style_name, (size, before, after) in specs.items():
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("透明房产小程序｜商业、办公功能需求文档")
    set_run_font(hr, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("透明房产小程序V1.1.0")
    set_run_font(r, size=23, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("商业、办公功能需求文档")
    set_run_font(r, size=15, bold=True, color=BLUE)

    meta = [
        ("版本", "V1.1.0"),
        ("需求时间", "2026/08/03"),
        ("需求", "周世杰（沿用 V1.0.0 模板）"),
        ("评审时间", "待评审"),
        ("系统", "小程序 / 移动端高保真原型"),
        ("需求类型", "V1.0.0 基础上的增量功能"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}：")
        set_run_font(r, size=10, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(9)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), MID_GRAY)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    add_para(doc, "修订记录", size=12.5, bold=True, after=6, keep=True)
    add_table(
        doc,
        ["日期", "版本", "调整/新增内容", "负责人"],
        [["2026/08/03", "V1.1.0", "新增商业、办公业态列表、筛选、详情、一房一价、配套和个人收藏链路。", "周世杰"]],
        [1400, 1200, 5260, 1500],
        font_size=8.7,
    )

    add_para(doc, "需求清单（仅调整和新增功能）", size=12.5, bold=True, before=3, after=6, keep=True)
    add_table(
        doc,
        ["端", "一级功能", "二级功能", "功能说明", "研发时间"],
        [
            ["C端", "首页", "楼盘业态", "增加住宅、商业、办公切换；文案统一为“住宅 / 商业 / 办公”。", "待排期"],
            ["C端", "首页", "筛选与列表", "按商业、办公的面积和总价口径动态切换筛选项及列表字段。", "待排期"],
            ["C端", "商业地图", "商业项目地图", "商业列表提供独立地图入口，仅展示商业项目；办公本期不开放独立地图。", "待排期"],
            ["C端", "项目详情", "商业/办公详情", "复用住宅信息架构，增加业态字段、销售情况、一房一价与周边配套。", "待排期"],
            ["C端", "项目详情", "一房一价", "商业、办公项目支持代表房源预览和完整一房一价列表。", "待排期"],
            ["C端", "我的", "我的关注/足迹", "改为子页面，支持住宅、商业、办公切换及详情跳转。", "待排期"],
            ["B端", "运营管理", "字段与数据接口", "本期无新增 B 端页面；补充业态、房源、配套及收藏/足迹的数据要求。", "待排期"],
        ],
        [700, 1300, 1750, 4510, 1100],
        font_size=8.2,
    )
    add_para(doc, "原型地址", size=11, bold=True, before=2, after=2, keep=True)
    add_para(doc, "线上基线：https://zhoushijie000.github.io/transparentHouse/index.html", size=9.5, color=BLUE, after=2)
    add_para(doc, f"当前本地原型：{ROOT / 'index.html'}?tab=commercial", size=9.5, color=MUTED, after=0)


def add_summary(doc):
    add_heading(doc, "0、需求概述", 1)
    add_callout(
        doc,
        "本期目标",
        "在不改变住宅楼盘主链路的前提下，新增商业（含公寓/商铺）与办公两类业态，使用户能够按业态完成“列表筛选 → 项目详情 → 一房一价 → 关注/足迹”的闭环浏览。",
    )
    add_table(
        doc,
        ["项", "说明"],
        [
            ["业务目标", "补齐非住宅项目展示能力，统一三类业态的基础浏览与咨询链路。"],
            ["用户价值", "公寓等低总价商业产品使用更贴近实际的价格段；办公项目使用独立面积、总价和办公类型筛选。"],
            ["范围内", "首页业态切换、动态筛选、商业地图、商业/办公详情、销售情况、一房一价、周边配套、关注与足迹子页面。"],
            ["范围外", "真实登录、接口联调、办公独立地图、B 端运营页面、在线交易、地图 SDK、跨设备收藏同步。"],
            ["兼容原则", "住宅既有字段、筛选、项目详情及板块地图保持原逻辑；新增能力不得影响住宅数据展示。"],
        ],
        [1900, 7460],
        font_size=9.2,
    )
    add_heading(doc, "0.1、信息架构与页面路由", 2)
    add_table(
        doc,
        ["入口", "页面/路由", "商业", "办公", "说明"],
        [
            ["楼盘", "index.html?tab={category}", "支持", "支持", "category：residential / commercial / office。"],
            ["地图", "commercial-map.html", "支持", "不支持", "商业独立项目地图；办公入口本期隐藏。"],
            ["详情", "business-project-detail.html?projectId={id}", "支持", "支持", "商业、办公共用详情壳层，数据按 category 渲染。"],
            ["一房一价", "project-one-price.html?projectId={id}", "支持", "支持", "按项目 id 加载预售证、楼栋、楼层和房源。"],
            ["我的关注", "my-follow.html?category={category}", "支持", "支持", "同时支持住宅；无筛选和地图入口。"],
            ["我的足迹", "my-history.html?category={category}", "支持", "支持", "同时支持住宅；无筛选和地图入口。"],
        ],
        [1200, 3000, 900, 900, 3360],
        font_size=8.6,
    )
    add_heading(doc, "0.2、核心链路", 2)
    for text in [
        "浏览链路：首页切换“商业/办公” → 搜索或筛选 → 点击楼盘卡片 → 进入商业/办公项目详情。",
        "房价链路：项目详情“一房一价” → 查看最低/最高代表房源 → 点击“查看完整一房一价” → 按预售证、楼栋和楼层查看房源。",
        "留存链路：列表或详情点击关注 → 我的关注切换对应业态 → 点击卡片返回项目详情。",
        "足迹链路：进入项目详情后自动记录 → 我的足迹切换对应业态 → 点击卡片返回项目详情。",
    ]:
        add_bullet(doc, text)


def add_home_requirements(doc):
    add_heading(doc, "1、C端", 1)
    add_heading(doc, "1.1、首页—楼盘业态与筛选", 2)
    add_para(doc, "首页楼盘列表在住宅基础上新增商业、办公两类业态。三项位于同一标签栏，切换时同步更新标题、项目数据、筛选字段、列表卡片和地图入口。", after=7)

    add_heading(doc, "（1）楼盘业态", 3)
    add_table(
        doc,
        ["展示项", "规则", "交互结果", "备注"],
        [
            ["住宅", "默认选中；沿用原住宅列表。", "URL 写入 tab=residential。", "不属于本期新增。"],
            ["商业", "展示商业（含公寓/商铺）项目。", "URL 写入 tab=commercial；刷新后保持。", "界面文案只展示“商业”。"],
            ["办公", "展示办公项目。", "URL 写入 tab=office；刷新后保持。", "界面文案只展示“办公”。"],
            ["切换重置", "切换业态后清空区域、板块、类型、面积、总价条件。", "立即刷新结果数与筛选按钮文案。", "避免不同业态条件串用。"],
        ],
        [1400, 2920, 2840, 2200],
        font_size=8.8,
    )

    add_heading(doc, "（2）筛选项", 3)
    add_table(
        doc,
        ["业态", "区域/板块", "第三筛选项", "面积", "总价"],
        [
            ["住宅", "保留", "户型：保留", "住宅面积：90㎡以下、90-120㎡、120-144㎡、144㎡以上", "住宅总价：沿用现状"],
            ["商业", "保留", "去掉产品类型", "文案“面积”：40㎡以下、40-60㎡、60-80㎡、80㎡以上", "文案“总价”：30万以内、30万-60万、60万-100万、100万以上"],
            ["办公", "保留", "办公类型：OFFICE / SOHO / LOFT 等", "办公面积：100㎡以下、100-150㎡、150-300㎡、300-500㎡、500㎡以上", "办公总价：300万以内、300-500万、500-800万、800-1000万、1000万及以上"],
        ],
        [1000, 1200, 1900, 2660, 2600],
        font_size=8.4,
    )
    add_callout(
        doc,
        "商业价格区间口径",
        "为适配百万以内公寓/商业产品，边界按“≤30万、>30且≤60万、>60且≤100万、>100万”理解；筛选命中采用项目总价区间与所选区间有交集即展示。",
        fill="FFF7E6",
        accent="A56500",
    )

    add_heading(doc, "（3）搜索与筛选交互", 3)
    add_table(
        doc,
        ["元素", "触发", "处理", "页面反馈"],
        [
            ["关键词", "输入楼盘名、区域、板块、标签等", "仅在当前业态数据中模糊匹配。", "列表与“共 X 个在售楼盘”实时更新。"],
            ["筛选按钮", "点击", "打开底部筛选弹层。", "有生效条件时显示“筛选 · N”。"],
            ["下拉选择", "任一条件变化", "按区域、板块、类型、面积、总价组合过滤。", "“查看 X 个楼盘”同步更新。"],
            ["重置", "点击重置", "清空当前业态全部筛选条件。", "恢复当前业态全部项目。"],
            ["空结果", "无项目命中", "列表为空。", "展示调整搜索或筛选条件的提示。"],
        ],
        [1500, 2100, 3280, 2480],
        font_size=8.8,
    )

    add_heading(doc, "1.2、首页—楼盘列表", 2)
    add_para(doc, "商业、办公列表沿用住宅楼盘卡片样式，保留相册横滑、关注、项目名称、位置、总价、标签与详情入口；底部四项信息按业态替换。")
    add_table(
        doc,
        ["字段区", "商业", "办公", "规则"],
        [
            ["价格", "参考总价区间（万元）", "参考总价区间（万元）", "显示 ¥最低万 - 最高万。"],
            ["标签", "集中式商业/专业市场/社区底商、装修等", "OFFICE/SOHO/LOFT、装修等", "最多按页面宽度换行展示。"],
            ["信息1", "产权年限", "产权年限", "缺失显示“—”。"],
            ["信息2", "商业建筑面积", "办公建筑面积", "单位 m²。"],
            ["信息3", "标准层面积", "层高", "按业态字段映射。"],
            ["信息4", "物业费", "物业费", "单位元/m²/月。"],
            ["卡片点击", "商业项目详情", "办公项目详情", "跳转 business-project-detail.html?projectId={id}。"],
            ["关注", "添加/取消关注", "添加/取消关注", "拦截整卡跳转，更新本地关注集合。"],
        ],
        [1700, 2700, 2700, 2260],
        font_size=8.6,
    )

    add_heading(doc, "1.3、地图入口", 2)
    add_table(
        doc,
        ["业态", "入口展示", "目标页面", "规则"],
        [
            ["住宅", "展示“住宅地图”", "chengdu-sector-map.html?mode=residential", "沿用原板块地图。"],
            ["商业", "展示“商业地图”", "commercial-map.html", "地图只展示商业项目，不显示住宅板块。"],
            ["办公", "本期隐藏", "—", "办公独立地图未实现，不展示空入口。"],
        ],
        [1200, 1800, 3260, 3100],
        font_size=9,
    )


def add_map_detail_requirements(doc):
    add_heading(doc, "1.4、商业地图", 2)
    add_para(doc, "商业列表页保留独立“商业地图”悬浮入口。地图以项目点位为主，不复用住宅板块边界和分级颜色。")
    add_table(
        doc,
        ["元素", "展示内容", "交互效果", "异常/状态"],
        [
            ["顶部", "返回、标题“商业地图 · 商业项目”", "返回商业列表 tab=commercial。", "始终展示。"],
            ["项目点位", "项目简称与定位点", "点击后点位高亮。", "默认选中首个项目。"],
            ["项目摘要", "项目名、区域位置、标签、起价", "随点位选择更新。", "字段缺失显示“—”。"],
            ["底部导航", "楼盘、数据、我的", "跳转对应主页面。", "楼盘项高亮。"],
        ],
        [1500, 2600, 3100, 2160],
        font_size=8.8,
    )
    add_callout(doc, "本期边界", "当前商业地图仅完成项目点位与摘要切换，不包含搜索、筛选、板块放大、聚合点位和办公项目地图。", fill=LIGHT_GRAY, accent=MUTED)

    add_heading(doc, "1.5、商业/办公项目详情", 2)
    add_para(doc, "商业和办公共用一套详情页面结构，视觉与住宅详情保持一致；业态差异字段集中到“基本信息”，不单独设置“物业服务”模块。")
    add_table(
        doc,
        ["模块", "商业", "办公", "交互/跳转"],
        [
            ["顶部概览", "相册、在售状态、项目名、区域、标签、参考总价", "同左，标签按办公类型展示", "相册可切换效果图/总平图/平面图/区位图/实景图；支持关注。"],
            ["快捷入口", "项目信息、销售信息、一房一价、周边配套、项目相册", "同左", "点击滚动到对应模块或打开相册。"],
            ["页签", "基本信息、销售信息、一房一价、周边配套", "同左", "图片离开视口后显示吸顶页签。"],
            ["基本信息", "商业字段 + 销售字段 + 物业公司/物业费", "办公字段 + 销售字段 + 物业公司/物业费/客梯数/货梯数", "超过 7 项时支持展开/收起。"],
            ["联系", "二维码、联系开发商", "二维码、联系开发商", "二维码按当前 URL 生成；拨打项目电话。"],
        ],
        [1500, 2500, 2500, 2860],
        font_size=8.4,
    )

    add_heading(doc, "（1）基本信息字段", 3)
    add_table(
        doc,
        ["字段组", "商业字段", "办公字段", "调整规则"],
        [
            ["项目基础", "项目名称、所属区域、所属板块、产权年限、商业建筑面积、标准层面积、户型区间、总户数", "项目名称、所属区域、所属板块、产权年限、办公建筑面积、层高、户型区间、总户数", "按 category 映射标签。"],
            ["销售信息", "开盘时间、交房时间、销售状态、营销代理", "同左", "合并展示在基本信息中。"],
            ["物业信息", "物业公司、物业费", "物业公司、物业费、客梯数、货梯数", "取消独立“物业服务”模块，字段合并到基本信息。"],
            ["移除字段", "服务时间、商业运营", "—", "商业详情不展示、前端数据结构不要求返回。"],
        ],
        [1500, 3100, 3100, 1660],
        font_size=8.2,
    )

    add_heading(doc, "（2）销售情况", 3)
    add_table(
        doc,
        ["字段", "计算/展示规则", "单位", "备注"],
        [
            ["总套数", "项目可售房源总量 total。", "套", "与住宅详情字段结构一致。"],
            ["已售套数", "sold。", "套", "不得大于总套数。"],
            ["在售套数", "max(0, total - sold)。", "套", "前端可计算，正式接口建议直接返回并校验。"],
            ["已售比例", "sold / total × 100%，保留 1 位小数。", "%", "total=0 时显示 0%。"],
            ["销售图", "环形图：在售蓝色、已售绿色。", "—", "图例同步展示套数。"],
            ["预计清盘时间", "展示预计清盘日期，仅作参考。", "年月", "生产数据需独立字段；不得复用交房时间。"],
        ],
        [1700, 3780, 900, 2980],
        font_size=8.7,
    )

    add_heading(doc, "（3）周边配套", 3)
    add_table(
        doc,
        ["配套分类", "商业示例", "办公示例", "展示规则"],
        [
            ["交通", "地铁站、公交站", "地铁站、交通枢纽", "默认选中；地图点位与列表同步。"],
            ["商业", "购物中心、社区商业", "商务商圈、购物中心", "展示名称与简述。"],
            ["公园", "桂溪生态公园、锦城湖公园", "交子公园、桂溪生态公园", "本期新增分类。"],
            ["服务", "医院、生活服务", "会议中心、金融机构", "按业态配置内容。"],
        ],
        [1200, 2500, 2500, 3160],
        font_size=8.8,
    )


def add_one_price_and_my(doc):
    add_heading(doc, "1.6、一房一价", 2)
    add_para(doc, "商业、办公项目新增一房一价能力，沿用住宅完整页面的预售证、楼栋/单元、楼层和房源结构。")
    add_table(
        doc,
        ["页面", "展示内容", "商业规则", "办公规则"],
        [
            ["详情预览", "总价最低、总价最高各 1 套代表房源", "房号使用栋/层/房号；副文案只展示面积，如“38㎡”。", "房号使用楼座/层/房号；副文案只展示面积。"],
            ["完整一房一价", "项目摘要、预售证页签、楼栋横滑、单元、楼层、房源卡", "单元文案为商铺，楼层按商业楼层生成。", "单元文案为 A区/B区，楼层按办公楼层生成。"],
            ["房源卡", "房号、状态、面积、总价", "不展示商业类型作为副文案。", "不展示办公类型作为副文案。"],
            ["返回", "返回项目详情", "回到 business-project-detail.html 并保持 projectId。", "同左。"],
        ],
        [1700, 2720, 2470, 2470],
        font_size=8.5,
    )
    add_heading(doc, "一房一价房源字段", 3)
    add_table(
        doc,
        ["字段", "定义", "类型", "必填", "备注"],
        [
            ["projectId", "项目唯一标识", "string", "是", "用于详情和一房一价关联。"],
            ["presaleId / presaleName", "预售证唯一标识/名称", "string", "是", "页签切换依据。"],
            ["building / unit / floor", "楼栋、单元/分区、楼层", "string/int", "是", "支持商业、办公差异化命名。"],
            ["roomNo", "房号", "string", "是", "页面主标题。"],
            ["area", "建筑面积", "number", "是", "单位㎡，展示时最多保留两位小数。"],
            ["totalPrice", "备案/参考总价", "number", "是", "单位万元。"],
            ["status", "在售/已售/其他", "enum", "是", "用于状态标签与筛选。"],
        ],
        [1500, 2850, 1200, 900, 2910],
        font_size=8.6,
    )

    add_heading(doc, "1.7、我的关注与我的足迹", 2)
    add_para(doc, "“我的关注”和“我的足迹”由我的页面入口进入独立子页面，卡片展示和楼盘列表一致，但不提供搜索、筛选和地图入口。")
    add_table(
        doc,
        ["功能", "我的关注", "我的足迹", "共同规则"],
        [
            ["业态切换", "住宅、商业、办公", "住宅、商业、办公", "默认住宅；URL category 参数可保持当前业态。"],
            ["统计文案", "已关注 X 个{业态}项目", "最近浏览 X 个{业态}项目", "只统计当前业态。"],
            ["列表卡片", "图片、名称、位置、总价、标签、面积、状态", "同左", "点击整卡进入对应详情。"],
            ["管理操作", "取消关注", "删除单条、清空全部足迹", "操作后立即刷新数量和空状态。"],
            ["空状态", "提示可在列表或详情页添加关注", "提示进入项目详情后自动记录", "按当前业态展示文案。"],
            ["限制", "无筛选、无地图", "无筛选、无地图", "与楼盘列表保持展示样式，不复制查找工具。"],
        ],
        [1400, 2500, 2500, 2960],
        font_size=8.5,
    )
    add_heading(doc, "数据与跳转规则", 3)
    for text in [
        "关注集合：透明房产原型使用 localStorage 键 transparentHouse:followedProjects，列表页和详情页均可写入。",
        "足迹集合：进入住宅、商业或办公详情时写入 transparentHouse:browsingHistory；再次访问时更新排序。",
        "商业/办公卡片跳转 business-project-detail.html?projectId={id}；住宅卡片跳转 project-detail.html?projectId={id}。",
        "本期原型为单设备本地存储，不包含登录态、云端同步、数据迁移和多端合并。",
    ]:
        add_bullet(doc, text)


def add_fields_states(doc):
    add_heading(doc, "1.8、项目字段定义", 2)
    add_table(
        doc,
        ["字段", "定义", "商业", "办公", "备注"],
        [
            ["projectId", "项目唯一标识", "commercial-*", "office-*", "路由与收藏主键。"],
            ["category", "业态编码", "commercial", "office", "前端分流必填。"],
            ["categoryLabel", "业态名称", "商业", "办公", "用户可见文案。"],
            ["name", "项目名称", "文本", "文本", "必填。"],
            ["district / sector", "区域/板块", "文本", "文本", "用于展示和筛选。"],
            ["address", "项目地址", "文本", "文本", "详情页展示。"],
            ["status", "项目销售状态", "在售/现房等", "在售/现房等", "徽标文案。"],
            ["priceMin / priceMax", "总价下限/上限", "万元", "万元", "用于列表和筛选。"],
            ["areaMin / areaMax", "面积下限/上限", "㎡", "㎡", "用于筛选；当前原型从面积区间文本解析。"],
            ["propertyYears", "产权年限", "年", "年", "常见为40年。"],
            ["commercialArea", "商业建筑面积", "适用", "不适用", "区间文本。"],
            ["standardFloorArea", "标准层面积", "适用", "可选", "商业卡片/详情。"],
            ["propertyLevel", "办公类型", "不适用", "OFFICE/SOHO/LOFT", "办公第三筛选项。"],
            ["officeArea", "办公建筑面积", "不适用", "适用", "区间文本。"],
            ["floorHeight", "层高", "可选", "适用", "单位m。"],
            ["propertyCompany / propertyFee", "物业公司/物业费", "适用", "适用", "合并到基本信息。"],
            ["passengerElevators / freightElevators", "客梯/货梯数量", "可选", "适用", "办公基础信息。"],
            ["total / sold", "总套数/已售套数", "适用", "适用", "驱动销售情况。"],
            ["tags / media / intro", "标签、相册、项目介绍", "适用", "适用", "与住宅同类字段。"],
        ],
        [1460, 2200, 1700, 1700, 2300],
        font_size=7.8,
    )

    add_heading(doc, "1.9、配套字段定义", 2)
    add_table(
        doc,
        ["字段", "定义", "类型", "必填", "备注"],
        [
            ["type", "配套分类", "enum", "是", "traffic / commerce / park / service。"],
            ["title", "配套名称", "string", "是", "列表标题与地图点位。"],
            ["desc", "距离或服务说明", "string", "是", "不强制固定距离格式。"],
            ["icon", "点位短标", "string", "否", "建议 1 个汉字。"],
            ["x / y", "原型地图相对坐标", "number", "原型必填", "正式地图接入后替换为经纬度。"],
        ],
        [1500, 2700, 1200, 1100, 2860],
        font_size=8.8,
    )

    add_heading(doc, "1.10、全局状态", 2)
    add_table(
        doc,
        ["状态", "触发", "页面表现", "可用操作"],
        [
            ["默认", "进入页面/合法 projectId", "展示当前业态默认数据和选中态。", "搜索、筛选、切换、进入详情。"],
            ["筛选后", "关键词或条件生效", "结果数、列表、按钮文案同步。", "继续调整或重置。"],
            ["空结果", "无项目匹配", "展示空状态。", "调整条件或重置。"],
            ["详情错误", "projectId 不存在", "提示未找到对应商业或办公项目。", "返回楼盘列表。"],
            ["无关注/足迹", "当前业态集合为空", "展示对应业态空状态。", "返回楼盘浏览。"],
            ["外链能力失败", "二维码服务/图片网络不可用", "保留基础文本与页面结构。", "复制地址或稍后重试。"],
        ],
        [1500, 2200, 3380, 2280],
        font_size=8.7,
    )


def add_acceptance(doc):
    add_heading(doc, "1.11、验收标准", 2)
    add_table(
        doc,
        ["编号", "验收场景", "前置条件", "预期结果"],
        [
            ["AC-01", "首页业态切换", "进入 index.html", "仅展示“住宅 / 商业 / 办公”；切换后列表、筛选、标题同步，URL tab 参数正确。"],
            ["AC-02", "商业筛选", "tab=commercial", "不显示产品类型；面积和总价文案为通用文案；价格段严格为30万以内、30-60万、60-100万、100万以上。"],
            ["AC-03", "办公筛选", "tab=office", "显示办公类型、办公面积、办公总价；选择条件后结果正确。"],
            ["AC-04", "列表字段", "商业/办公有数据", "商业显示产权、商业面积、标准层面积、物业费；办公显示产权、办公面积、层高、物业费。"],
            ["AC-05", "详情跳转", "点击商业/办公卡片", "进入 business-project-detail.html，projectId 与卡片一致。"],
            ["AC-06", "商业地图", "tab=commercial", "显示商业地图入口；进入后只展示商业点位；点击点位更新摘要。"],
            ["AC-07", "办公地图边界", "tab=office", "不展示地图入口，不出现无效空链接。"],
            ["AC-08", "详情信息架构", "进入商业/办公详情", "无独立物业服务模块；销售字段和物业字段合并到基本信息；商业不展示服务时间、商业运营。"],
            ["AC-09", "销售情况", "total、sold 合法", "总套数、已售、在售、已售比例和环图结果一致；在售不得为负数。"],
            ["AC-10", "一房一价预览", "进入详情价格区", "展示最低/最高两套；房号下副文案只显示面积，不显示产品类型或状态。"],
            ["AC-11", "完整一房一价", "点击查看更多", "按项目加载商业/办公房源；返回时回到同一项目详情。"],
            ["AC-12", "周边公园", "切换周边配套", "商业、办公均有“公园”分类，地图点位与列表内容同步。"],
            ["AC-13", "我的关注", "已关注至少一个不同业态项目", "子页面可切换三业态；点击卡片进详情；可取消关注；无筛选和地图。"],
            ["AC-14", "我的足迹", "浏览过至少一个不同业态详情", "子页面可切换三业态；点击卡片进详情；可删单条/清空；无筛选和地图。"],
            ["AC-15", "住宅回归", "切换回住宅", "住宅列表、筛选、地图和详情链路保持原功能。"],
        ],
        [800, 1900, 2200, 4460],
        font_size=7.9,
    )

    add_heading(doc, "1.12、测试重点", 2)
    for text in [
        "边界值：30、60、100、300、500、800、1000 万，以及 40、60、80、100、150、300、500㎡。",
        "业态切换后筛选残留、URL 参数回显、刷新恢复、浏览器后退。",
        "列表关注按钮与整卡跳转事件冲突；关注状态在列表、详情、我的关注三处一致。",
        "非法 projectId、空数据、total=0、sold>total、缺失图片/电话/配套时的兜底。",
        "商业/办公完整一房一价返回链接保持项目 id，不误回住宅详情。",
    ]:
        add_bullet(doc, text)


def add_backend_and_release(doc):
    add_heading(doc, "2、B端及数据影响（本期暂无新增页面）", 1)
    add_para(doc, "当前 V1.1.0 为 C 端原型增量。为后续正式接入，B 端或接口层需具备以下能力；若本期仅交付原型，可作为后续版本待办。")
    add_table(
        doc,
        ["能力", "新增/调整项", "校验规则", "优先级"],
        [
            ["项目管理", "增加 category：住宅/商业/办公；按业态显示字段。", "projectId 唯一；category 必填且不可随意变更。", "P0"],
            ["价格与面积", "存储数值下限/上限，前端负责文案。", "min≤max；价格单位万元、面积单位㎡。", "P0"],
            ["商业字段", "商业建筑面积、标准层面积、产权年限、物业公司、物业费。", "服务时间、商业运营不在本期字段清单。", "P0"],
            ["办公字段", "办公类型、办公建筑面积、层高、客梯数、货梯数、产权与物业字段。", "办公类型需配置字典。", "P0"],
            ["销售数据", "总套数、已售套数、在售套数、已售比例、预计清盘时间。", "已售≤总套数；预计清盘时间独立于交房时间。", "P0"],
            ["一房一价", "预售证、楼栋、单元/分区、楼层、房号、面积、总价、状态。", "项目、预售证与房源层级关联完整。", "P0"],
            ["配套管理", "增加 park 分类；商业/办公可配置不同配套。", "名称必填；正式地图使用经纬度。", "P1"],
            ["关注/足迹", "如接入登录，提供查询、新增、删除、清空接口。", "幂等；按 category 分页；支持时间排序。", "P1"],
            ["商业地图", "项目点位、名称、区域、标签、起价。", "经纬度合法；办公地图本期不配置入口。", "P1"],
        ],
        [1700, 3480, 3000, 1180],
        font_size=8.2,
    )

    add_heading(doc, "2.1、建议接口返回原则", 2)
    for text in [
        "列表接口按 category 返回对应业态项目，并提供可筛选的数值字段，避免前端解析展示文案。",
        "详情接口将 basic、sales、property、media、support 分组返回；前端按 category 配置标签。",
        "一房一价接口按 projectId → presale → building → unit → floor → rooms 组织，支持仅在售查询。",
        "所有金额字段明确单位；缺失值返回 null，不使用“暂无/待定”等展示文案作为数据值。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "3、待确认事项与风险", 1)
    add_table(
        doc,
        ["编号", "事项", "当前处理", "需要确认"],
        [
            ["Q1", "“商业”是否长期同时覆盖公寓、商铺及集中式商业", "原型统一使用 commercial 分类，界面显示“商业”。", "产品分类与后台字典。"],
            ["Q2", "办公独立地图", "本期隐藏入口。", "后续是否建设、是否复用商业项目地图。"],
            ["Q3", "预计清盘时间来源", "原型有参考展示。", "正式数据源和更新频率。"],
            ["Q4", "关注/足迹登录同步", "原型使用 localStorage。", "是否接入账号体系及跨端同步。"],
            ["Q5", "商业地图详情入口", "当前点位仅切换摘要。", "是否增加“查看详情”按钮或点击摘要整卡跳转。"],
            ["Q6", "商业/办公筛选区间是否运营可配", "原型写死区间。", "后端配置化或版本内固定。"],
        ],
        [800, 2850, 3400, 2310],
        font_size=8.5,
    )
    add_callout(doc, "主要风险", "当前原型中的商业/办公数据为演示数据，正式上线前必须完成字段映射、价格/面积单位校验、房源状态口径统一和真实图片/地图数据接入。", fill="FFF1F1", accent=RED)

    add_heading(doc, "4、发布与回归建议", 1)
    add_table(
        doc,
        ["阶段", "内容", "完成标准"],
        [
            ["开发自测", "三业态切换、筛选边界、详情、一房一价、关注/足迹。", "AC-01 至 AC-14 全部通过。"],
            ["数据验收", "商业/办公字段、单位、价格区间、销售套数、配套类型。", "抽检项目与后台/接口一致，无负数、空主键或单位混用。"],
            ["住宅回归", "住宅列表、筛选、地图、详情、一房一价。", "AC-15 通过，无样式或路由回归。"],
            ["灰度发布", "先开放商业与办公入口，监控空结果、详情错误及联系方式点击。", "核心链路无阻断；异常可回退隐藏新增业态入口。"],
        ],
        [1500, 4100, 3760],
        font_size=8.8,
    )
    add_para(doc, "—— 文档结束 ——", size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=0)


def audit_document(doc):
    section = doc.sections[0]
    assert abs(int(section.page_width) - int(Cm(21))) <= 5000
    assert abs(int(section.page_height) - int(Cm(29.7))) <= 5000
    assert len(doc.tables) >= 15
    for table in doc.tables:
        assert table.autofit is False
        grid = table._tbl.tblGrid
        widths = [int(col.get(qn("w:w"))) for col in grid]
        assert sum(widths) == CONTENT_WIDTH_DXA, widths


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    doc.core_properties.title = "透明房产小程序V1.1.0—商业、办公功能需求文档"
    doc.core_properties.subject = "透明房产小程序在 V1.0.0 基础上增加商业、办公功能的增量 PRD"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "透明房产,商业,办公,PRD,V1.1.0"

    add_cover(doc)
    add_summary(doc)
    add_home_requirements(doc)
    add_map_detail_requirements(doc)
    add_one_price_and_my(doc)
    add_fields_states(doc)
    add_acceptance(doc)
    add_backend_and_release(doc)

    audit_document(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
