from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\ZhuanZ\Desktop\zhou\HouseFuni\透明房产网小程序项目详情二维码.docx")
IMG = Path(r"C:\Users\ZhuanZ\AppData\Local\Temp\codex-clipboard-9dcf833d-3f12-4796-b895-a90438367378.png")
BLUE, NAVY, LIGHT, GRAY = "2E74B5", "0B2545", "E8EEF5", "667085"

def font(run, name="Calibri", size=10.5, color=None, bold=False):
    run.font.name=name
    rpr=run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name); rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size=Pt(size); run.bold=bold
    if color: run.font.color.rgb=RGBColor.from_string(color)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr()
    shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for edge,val in [("top",80),("start",120),("bottom",80),("end",120)]:
        node=tcMar.find(qn("w:"+edge))
        if node is None: node=OxmlElement("w:"+edge); tcMar.append(node)
        node.set(qn("w:w"),str(val)); node.set(qn("w:type"),"dxa")

def table_setup(table, widths):
    table.autofit=False
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col=OxmlElement("w:gridCol"); col.set(qn("w:w"),str(w)); grid.append(col)
    for row in table.rows:
        for cell,w in zip(row.cells,widths):
            tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.first_child_found_in("w:tcW")
            if tcW is None: tcW=OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"),str(w)); tcW.set(qn("w:type"),"dxa"); margins(cell)

def borders(table):
    tblPr=table._tbl.tblPr; bs=tblPr.first_child_found_in("w:tblBorders")
    if bs is None: bs=OxmlElement("w:tblBorders"); tblPr.append(bs)
    for edge in ("top","left","bottom","right","insideH","insideV"):
        node=OxmlElement("w:"+edge); node.set(qn("w:val"),"single"); node.set(qn("w:sz"),"6"); node.set(qn("w:color"),"D0D5DD"); bs.append(node)

def code(doc,text):
    p=doc.add_paragraph(style="Code")
    for i,line in enumerate(text.splitlines()):
        r=p.add_run(line); font(r,"Consolas",9,NAVY)
        if i<len(text.splitlines())-1: r.add_break()

def bullet(doc,text):
    doc.add_paragraph(text, style="List Bullet")

doc=Document()
sec=doc.sections[0]
for attr in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(sec,attr,Inches(0.8 if attr in ("top_margin","bottom_margin") else 0.85))
sec.header_distance=Inches(0.35); sec.footer_distance=Inches(0.35)
normal=doc.styles["Normal"]; normal.font.name="Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri"); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string("27364B"); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for nm,sz,col,bef,aft in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7)]:
    s=doc.styles[nm]; s.font.name="Calibri"; s._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); s._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri"); s.font.size=Pt(sz); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(bef); s.paragraph_format.space_after=Pt(aft); s.paragraph_format.keep_with_next=True
cs=doc.styles.add_style("Code",WD_STYLE_TYPE.PARAGRAPH); cs.font.name="Consolas"; cs._element.rPr.rFonts.set(qn("w:ascii"),"Consolas"); cs._element.rPr.rFonts.set(qn("w:hAnsi"),"Consolas"); cs.font.size=Pt(9); cs.paragraph_format.left_indent=Inches(0.12); cs.paragraph_format.right_indent=Inches(0.12); cs.paragraph_format.space_after=Pt(8)

p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); font(p.add_run("透明房产网小程序项目详情二维码"),"Calibri",25,NAVY,True)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(14); font(p.add_run("接口文档 · 生成项目详情二维码"),"Calibri",12,BLUE,True)

meta=doc.add_table(rows=2,cols=4); meta.alignment=WD_TABLE_ALIGNMENT.LEFT; table_setup(meta,[1450,3230,1450,3230]); borders(meta)
for i,(k,v) in enumerate([("接口名称","项目详情二维码生成"),("所属模块","小程序项目详情"),("请求方式","POST"),("版本状态","正式版默认")]):
    row=i//2; col=(i%2)*2; meta.cell(row,col).text=k; meta.cell(row,col+1).text=v; shade(meta.cell(row,col),LIGHT)
    for rr in meta.cell(row,col).paragraphs[0].runs: font(rr,size=10,color=NAVY,bold=True)
    for rr in meta.cell(row,col+1).paragraphs[0].runs: font(rr,size=10)

doc.add_heading("1. 接口概览",1)
doc.add_paragraph("用于生成透明房产网小程序项目详情页二维码。调用方通过传入小程序页面路径、项目分期标识及物业/房屋类型，获得对应的二维码图片或二维码数据。")
doc.add_heading("请求信息",2)
info=doc.add_table(rows=5,cols=2); table_setup(info,[2700,6660]); borders(info)
for i,(k,v) in enumerate([("请求地址","https://ibs.funi.com/zhfw/hcsp/t/xcx/detail/info/dimensionalCode"),("请求方法","POST"),("Content-Type","application/json"),("鉴权方式","x-inc-authorization: Bearer <token>"),("签名开关","skipSign: true")]):
    info.cell(i,0).text=k; info.cell(i,1).text=v; shade(info.cell(i,0),LIGHT)
    for rr in info.cell(i,0).paragraphs[0].runs: font(rr,size=10,color=NAVY,bold=True)
    for rr in info.cell(i,1).paragraphs[0].runs: font(rr,size=10)

doc.add_heading("2. 请求头参数",1)
t=doc.add_table(rows=1,cols=4); table_setup(t,[1900,1650,1200,4610]); borders(t)
for j,x in enumerate(["参数名","类型","是否必填","说明 / 示例"]): t.cell(0,j).text=x; shade(t.cell(0,j),LIGHT)
for row in [("skipSign","string","是","true；示例中关闭签名校验"),("x-inc-authorization","string","是","Bearer cb2500e6-a112-4c85-925b-127619b35089"),("Content-Type","string","是","application/json")]:
    cells=t.add_row().cells
    for j,x in enumerate(row): cells[j].text=x
for row in t.rows:
    for cell in row.cells:
        for p0 in cell.paragraphs:
            for rr in p0.runs: font(rr,size=9.5,color=NAVY if row is t.rows[0] else None,bold=row is t.rows[0])

doc.add_heading("3. 请求体参数",1)
t=doc.add_table(rows=1,cols=5); table_setup(t,[1800,1500,1200,2200,2660]); borders(t)
for j,x in enumerate(["参数名","类型","必填","默认值","说明"]): t.cell(0,j).text=x; shade(t.cell(0,j),LIGHT)
params=[("path","string","是","/subpackage/buildingDetail/pages/index","要打开的小程序页面路径"),("width","integer","否","0","二维码宽度；0 表示使用服务端默认值"),("version","string","否","release","小程序版本：release 正式版、trial 体验版、develop 开发版"),("text","string","否","string","需要生成的二维码文字"),("communityId","string","是","string","分期 Id"),("propertyType","string","否","HOUSE","物业类型；默认 HOUSE"),("houseType","string","否","1001","房屋类型：1001 住宅、1002 商业、1003 办公、1004 车位")]
for row in params:
    cells=t.add_row().cells
    for j,x in enumerate(row): cells[j].text=x
for row in t.rows:
    for cell in row.cells:
        for p0 in cell.paragraphs:
            for rr in p0.runs: font(rr,size=9.2,color=NAVY if row is t.rows[0] else None,bold=row is t.rows[0])

doc.add_heading("4. 调用示例",1)
doc.add_heading("cURL",2)
code(doc, """curl --location --request POST 'https://ibs.funi.com/zhfw/hcsp/t/xcx/detail/info/dimensionalCode' \\
--header 'skipSign: true' \\
--header 'x-inc-authorization: Bearer cb2500e6-a112-4c85-925b-127619b35089' \\
--header 'Content-Type: application/json' \\
--data-raw '{
  \"path\": \"/subpackage/buildingDetail/pages/index\",
  \"width\": 0,
  \"version\": \"string\",
  \"text\": \"string\",
  \"communityId\": \"string\",
  \"propertyType\": \"HOUSE\",
  \"houseType\": \"1001\"
}'""")
doc.add_heading("JSON 请求体",2)
code(doc, """{
  "path": "/subpackage/buildingDetail/pages/index",
  "width": 0,
  "version": "string",
  "text": "string",
  "communityId": "string",
  "propertyType": "HOUSE",
  "houseType": "1001"
}""")

doc.add_heading("5. 返回响应",1)
doc.add_paragraph("截图展示了“返回响应”区域，但未包含具体响应字段内容。因此本文仅确认接口用途为生成项目详情二维码，具体响应结构（例如二维码图片地址、Base64 数据、状态码及错误信息）应以实际接口返回或后端 Swagger 定义为准。")
t=doc.add_table(rows=1,cols=3); table_setup(t,[2200,2200,4960]); borders(t)
for j,x in enumerate(["场景","建议处理","备注"]): t.cell(0,j).text=x; shade(t.cell(0,j),LIGHT)
for row in [("HTTP 200","读取响应中的二维码结果","字段名需以后端实际返回为准"),("HTTP 4xx","提示参数或鉴权错误","检查 communityId、Token 及请求头"),("HTTP 5xx","提示服务暂不可用并记录 requestId","建议重试并联系接口维护方")]:
    cells=t.add_row().cells
    for j,x in enumerate(row): cells[j].text=x
for row in t.rows:
    for cell in row.cells:
        for p0 in cell.paragraphs:
            for rr in p0.runs: font(rr,size=9.5,color=NAVY if row is t.rows[0] else None,bold=row is t.rows[0])

doc.add_heading("6. 使用说明与注意事项",1)
for x in ["communityId 为必填业务标识，应传入需要生成二维码的项目分期 Id。","path 应保持为小程序实际可访问的页面路径：/subpackage/buildingDetail/pages/index。","propertyType 与 houseType 应保持业务枚举一致；商业场景建议传 propertyType=HOUSE、houseType=1002，并以服务端规则为准。","生产环境不要将长期有效的 Bearer Token 固化在前端、日志或公开文档中；本文保留示例 Token 仅用于还原接口示例。","version 未传时建议使用 release；联调阶段可根据环境选择 trial 或 develop。"]: bullet(doc,x)

if IMG.exists():
    doc.add_heading("附录 A. 参数页面截图依据",1)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(IMG),width=Inches(6.35))
    p=doc.add_paragraph("图 A-1  项目详情二维码接口 Body 参数及示例"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for rr in p.runs: font(rr,size=9,color=GRAY)

doc.core_properties.title="透明房产网小程序项目详情二维码"; doc.core_properties.subject="接口文档"; doc.core_properties.author="OpenAI Codex"
doc.save(OUT)
print(OUT)

